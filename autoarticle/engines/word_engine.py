from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import datetime

from .constants import THEMES
from .utils import get_valid_images


class WordEngine:
    def __init__(self, theme_name, school_name, layout_version="v1"):
        self.theme = THEMES.get(theme_name, THEMES[next(iter(THEMES))])
        self.school_name = school_name
        self.layout_version = layout_version

    def _cover_sizes(self):
        if self.layout_version == "v6":
            return 42, 18
        if self.layout_version == "v5":
            return 41, 19
        if self.layout_version == "v4":
            return 42, 18
        if self.layout_version in {"v2", "v3"}:
            return 40, 18
        return 44, 20

    def _meta_fill_hex(self):
        if self.layout_version == "v6":
            return "F8F8F8"
        if self.layout_version == "v5":
            return "EFE9DF"
        if self.layout_version == "v4":
            return "F0F0F0"
        if self.layout_version == "v3":
            return "EFEFE8"
        if self.layout_version == "v2":
            return "F2F6FC"
        return "F5F5F5"

    def generate(self, articles):
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Malgun Gothic"
        style.font.size = Pt(11)

        self._add_cover_page(doc)

        for art in articles:
            doc.add_page_break()
            self._add_article_page(doc, art)

        self._add_header_footer(doc)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _add_header_footer(self, doc):
        section = doc.sections[0]

        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = f"{self.school_name} 소식지"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*self.theme["main"])
            run.font.bold = True

        divider = header.add_paragraph()
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = divider.add_run("─" * 50)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(8)

        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = footer_para.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    def _add_cover_page(self, doc):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(180)

        title_size, subtitle_size = self._cover_sizes()

        title = doc.add_heading(level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(self.school_name)
        run.font.size = Pt(title_size)
        run.font.color.rgb = RGBColor(*self.theme["main"])
        run.font.bold = True

        now = datetime.datetime.now()
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"{now.year}년 {now.month}월 뉴스레터")
        run.font.size = Pt(subtitle_size)
        run.font.color.rgb = RGBColor(90, 90, 90 if self.layout_version != "v3" else 70)

        doc.add_paragraph()
        deco = doc.add_paragraph()
        deco.alignment = WD_ALIGN_PARAGRAPH.CENTER
        deco_run = deco.add_run("◆ ◆ ◆")
        deco_run.font.size = Pt(16)
        deco_run.font.color.rgb = RGBColor(*self.theme["accent"])

        doc.add_paragraph()
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"발행일: {now.strftime('%Y년 %m월 %d일')}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(120, 120, 120)

    def _add_article_page(self, doc, article):
        title = doc.add_heading(level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run(str(article.get("title", "")))
        run.font.size = Pt(20 if self.layout_version in {"v2", "v3", "v4", "v5", "v6"} else 22)
        run.font.color.rgb = RGBColor(*self.theme["main"])
        run.font.bold = True

        meta_table = doc.add_table(rows=1, cols=1)
        meta_table.style = "Light Grid Accent 1"
        cell = meta_table.cell(0, 0)

        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), self._meta_fill_hex())
        cell._element.get_or_add_tcPr().append(shading)

        info_parts = []
        if article.get("date"):
            info_parts.append(f"일시: {article['date']}")
        if article.get("location"):
            info_parts.append(f"장소: {article['location']}")
        if article.get("grade"):
            info_parts.append(f"대상: {article['grade']}")

        p = cell.paragraphs[0]
        meta_run = p.add_run("  |  ".join(info_parts))
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(80, 80, 80)

        doc.add_paragraph()

        valid_imgs = get_valid_images(article.get("images", []), max_count=4)
        if valid_imgs:
            count = len(valid_imgs)
            if count == 1:
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_cell = table.cell(0, 0)
                ip = img_cell.paragraphs[0]
                ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ir = ip.add_run()
                ir.add_picture(valid_imgs[0], width=Inches(5.0))

                tc_pr = img_cell._element.get_or_add_tcPr()
                borders = OxmlElement("w:tcBorders")
                for name in ["top", "left", "bottom", "right"]:
                    border = OxmlElement(f"w:{name}")
                    border.set(qn("w:val"), "single")
                    border.set(qn("w:sz"), "12")
                    border.set(qn("w:color"), "CCCCCC")
                    borders.append(border)
                tc_pr.append(borders)
            elif count == 2:
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.allow_autofit = False
                for idx in range(2):
                    img_cell = table.cell(0, idx)
                    ip = img_cell.paragraphs[0]
                    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    ir = ip.add_run()
                    ir.add_picture(valid_imgs[idx], width=Inches(3.2))
            else:
                table = doc.add_table(rows=2, cols=2)
                table.autofit = False
                table.allow_autofit = False
                for idx, img in enumerate(valid_imgs[:4]):
                    row = idx // 2
                    col = idx % 2
                    img_cell = table.cell(row, col)
                    ip = img_cell.paragraphs[0]
                    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    ir = ip.add_run()
                    ir.add_picture(img, width=Inches(2.8))

        doc.add_paragraph()

        body = doc.add_paragraph()
        body.paragraph_format.line_spacing = 1.5
        body.paragraph_format.space_after = Pt(12)
        run = body.add_run(str(article.get("content", "")))
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(40, 40, 40)

        doc.add_paragraph()
        sep = doc.add_paragraph()
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sep.add_run("◆ ◆ ◆")
        sr.font.color.rgb = RGBColor(*self.theme["accent"])
        sr.font.size = Pt(14)
